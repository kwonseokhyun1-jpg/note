#!/usr/bin/env python3
"""Generate AI agent insurance onboarding Word document."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(f"{bold_prefix}: ")
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def build_document(csv_path: Path, output_path: Path) -> None:
    rows = list(csv.DictReader(csv_path.open(encoding="utf-8")))

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("AI Agent Insurance Onboarding Guide", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("How to get an AI agent insured — provider flows, timelines, and requirements")
    run.italic = True

    doc.add_paragraph(
        "This document summarizes how companies can obtain insurance for AI agents in 2026. "
        "Onboarding varies sharply by provider: some require adversarial testing and certification "
        "before binding, others offer same-day quotes with no technical audit, and enterprise buyers "
        "typically route through brokers over weeks or months."
    )

    add_heading(doc, "Three onboarding archetypes", 1)

    archetypes = [
        (
            "1. Certification-first (agent-native MGAs)",
            "Klaimee, AIUC, Ollive (pre-bind)",
            "You must demonstrate agent safety through technical evaluation before insurance is offered. "
            "Best when enterprise procurement demands third-party proof that your agent is fit to deploy.",
        ),
        (
            "2. Scan-then-insure (security-integrated)",
            "Mount",
            "Vulnerability scanning and risk scoring precede or run alongside underwriting. "
            "Coverage pricing may adjust as your security posture improves.",
        ),
        (
            "3. Fast self-serve (startup insurance + AI module)",
            "Corgi, Vouch",
            "Standard digital application (5–10 minutes), instant or same-day quote, no mandatory red-team. "
            "AI-specific risks are covered via Tech E&O, AI Liability, or an AI endorsement.",
        ),
    ]
    for name, providers, desc in archetypes:
        p = doc.add_paragraph()
        p.add_run(name).bold = True
        p.add_run(f" — {providers}\n{desc}")

    add_heading(doc, "At-a-glance comparison", 1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Provider", "Model", "Timeline", "Technical audit required?"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "D9E2F3")

    audit_map = {
        "Klaimee": "Yes — 100+ adversarial probes",
        "Mount": "Yes — vulnerability scanning",
        "Corgi": "No",
        "Vouch": "No (optional AI endorsement)",
        "AIUC": "Yes — AIUC-1 certification",
        "Armilla AI": "Yes — risk assessment",
        "Testudo": "No",
        "Founder Shield": "No",
        "Ollive": "Yes — Agent Trust Score evals",
        "Luphra": "N/A (not available)",
        "Marsh / Aon / Gallagher": "Varies by carrier",
        "Kelly Insurance Group": "Depends on carrier",
    }

    for row in rows:
        provider = row["Provider"]
        if provider == "Luphra":
            continue
        cells = table.add_row().cells
        cells[0].text = provider
        cells[1].text = row["Onboarding_Model"]
        cells[2].text = row["Typical_Timeline"]
        cells[3].text = audit_map.get(provider, "Varies")

    doc.add_paragraph()

    add_heading(doc, "Provider-by-provider onboarding", 1)

    for row in rows:
        provider = row["Provider"]
        add_heading(doc, provider, 2)

        meta = doc.add_paragraph()
        meta.add_run("Type: ").bold = True
        meta.add_run(f"{row['Provider_Type']}  |  ")
        meta.add_run("Channel: ").bold = True
        meta.add_run(f"{row['Channel']}  |  ")
        meta.add_run("Timeline: ").bold = True
        meta.add_run(row["Typical_Timeline"])

        add_heading(doc, "Steps", 3)
        for i in range(1, 6):
            step = row.get(f"Step_{i}", "").strip()
            if step:
                add_bullet(doc, step, f"Step {i}")

        add_heading(doc, "What you need to provide", 3)
        doc.add_paragraph(row["Inputs_Required"])

        add_heading(doc, "Pre-bind requirements", 3)
        doc.add_paragraph(row["Pre_Bind_Requirements"])

        add_heading(doc, "What you receive", 3)
        doc.add_paragraph(row["Deliverables"])

        if row["Notes"]:
            add_heading(doc, "Notes", 3)
            doc.add_paragraph(row["Notes"])

        p = doc.add_paragraph()
        p.add_run("Website: ").bold = True
        p.add_run(row["Website"])

    add_heading(doc, "Broker and enterprise paths", 1)
    doc.add_paragraph(
        "Testudo and Founder Shield require an existing broker relationship — there is no direct "
        "self-serve application. Armilla operates through surplus-lines brokers for Lloyd's-backed coverage. "
        "Marsh, Aon, and Gallagher serve large enterprises with multi-week placement cycles, manuscript "
        "endorsements, and policy towers that may combine cyber, tech E&O, and standalone generative-AI liability."
    )

    add_heading(doc, "Choosing a path", 1)
    choices = [
        (
            "Closing an enterprise deal this week",
            "Klaimee (certification + guarantee pack for procurement) or Corgi/Vouch (same-day COI)",
        ),
        (
            "Selling AI agents to enterprises",
            "Klaimee or AIUC — certification is the product procurement teams want",
        ),
        (
            "Deploying gen-AI internally (not building it)",
            "Testudo via broker — built for deployers, not vendors; no stack audit",
        ),
        (
            "Need performance warranty in contract",
            "Armilla Guaranteed — separate from insurance; funds remedy if KPIs are missed",
        ),
        (
            "Not ready to bind yet",
            "Ollive — Agent Trust Score and risk evals prepare you for future AI E&O",
        ),
    ]
    for scenario, recommendation in choices:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(scenario).bold = True
        p.add_run(f" → {recommendation}")

    add_heading(doc, "Coverage gap context", 1)
    doc.add_paragraph(
        "Most existing cyber and tech E&O policies do not affirmatively cover autonomous AI agent actions. "
        "ISO/Verisk generative-AI exclusions (CG 40 47, CG 40 48, CG 35 08) are attaching at CGL renewals "
        "from January 2026. Purpose-built AI agent insurance fills the gap between cyber (external attacks) "
        "and tech E&O (human-written software bugs) for harm caused when an agent acts on its own."
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Research compiled July 2026. Verify current terms with each provider before binding.")
    run.italic = True
    run.font.size = Pt(9)

    doc.save(output_path)


if __name__ == "__main__":
    base = Path(__file__).resolve().parent
    csv_file = base / "ai-agent-insurance-onboarding.csv"
    out_file = base / "ai-agent-insurance-onboarding.docx"
    build_document(csv_file, out_file)
    print(f"Wrote {out_file}")
